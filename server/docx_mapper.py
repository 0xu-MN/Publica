"""
docx_mapper.py — Coordinate-based DOCX Template Mapping Engine

Two modes:
  1. COORDINATE MODE: If template is recognized (via template_registry), use exact
     cell coordinates to insert content. Maximum precision.
  2. HEURISTIC MODE: For unrecognized forms, fallback to keyword-scanning heuristics.
"""
import docx
import re
from bs4 import BeautifulSoup
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from template_registry import detect_template, normalize_section_key, get_unique_cells, analyze_template_with_ai


# ─────────────────────────────────────────────
# Shared Utilities
# ─────────────────────────────────────────────

def iter_block_items(parent):
    """Yield each paragraph and table child within *parent*, in document order."""
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("unsupported parent")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _reset_paragraph_format(p):
    """Forcefully cure overlapping text bugs in strict government forms."""
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def _apply_style(run, style_config):
    """Apply font style from template config to a run."""
    if not style_config:
        return
    if "font_name" in style_config:
        run.font.name = style_config["font_name"]
    if "font_size" in style_config:
        run.font.size = Pt(style_config["font_size"])


def __render_markdown_text(p, text, style_config=None):
    """Parse **bold** markers and handle newlines. Optionally apply style."""
    text = text.replace('\\n', '\n')
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith('**') and part.endswith('**')
        clean_text = part[2:-2] if is_bold else part
        
        lines = clean_text.split('\n')
        for i, line in enumerate(lines):
            if line:
                run = p.add_run(line)
                if is_bold:
                    run.bold = True
                _apply_style(run, style_config)
            if i < len(lines) - 1:
                p.add_run().add_break()


def _clear_cell(cell):
    """Remove all existing content from a cell, preserving cell structure."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.text = ""
    # Remove extra paragraphs beyond the first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)


def _insert_html_into_cell(cell, elements, style_config=None):
    """Insert parsed HTML elements into a specific cell with style preservation."""
    # Clear existing content first
    _clear_cell(cell)
    
    first_element = True
    for element in elements:
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(element.name[1])
            if first_element:
                p = cell.paragraphs[0]
                first_element = False
            else:
                p = cell.add_paragraph()
            _reset_paragraph_format(p)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(element.get_text().strip())
            run.bold = True
            size = min(12, max(9, 13 - level))
            run.font.size = Pt(size)
            _apply_style(run, style_config)

        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                if first_element:
                    p = cell.paragraphs[0]
                    first_element = False
                else:
                    p = cell.add_paragraph()
                _reset_paragraph_format(p)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                __render_markdown_text(p, text, style_config)

        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                if first_element:
                    p = cell.paragraphs[0]
                    first_element = False
                else:
                    p = cell.add_paragraph()
                _reset_paragraph_format(p)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.left_indent = Pt(10)
                bullet_run = p.add_run("• ")
                _apply_style(bullet_run, style_config)
                __render_markdown_text(p, li.get_text().strip(), style_config)

        elif element.name == 'table':
            # Skip nested tables inside cells for now — just extract text
            rows = element.find_all('tr')
            for r in rows:
                cells_html = r.find_all(['td', 'th'])
                row_text = " | ".join(c.get_text().strip() for c in cells_html)
                if row_text.strip():
                    if first_element:
                        p = cell.paragraphs[0]
                        first_element = False
                    else:
                        p = cell.add_paragraph()
                    _reset_paragraph_format(p)
                    __render_markdown_text(p, row_text, style_config)

        elif element.name is None:
            text = str(element).strip()
            if text:
                if first_element:
                    p = cell.paragraphs[0]
                    first_element = False
                else:
                    p = cell.add_paragraph()
                _reset_paragraph_format(p)
                __render_markdown_text(p, text, style_config)


def insert_elements_into_container(container, elements, style_config=None):
    """Insert elements into a Document or Cell container (heuristic mode)."""
    for element in elements:
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(element.name[1])
            p = container.add_paragraph()
            _reset_paragraph_format(p)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(element.get_text().strip())
            run.bold = True
            run.font.size = Pt(13 - (level * 1))

        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                p = container.add_paragraph()
                _reset_paragraph_format(p)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                __render_markdown_text(p, text, style_config)

        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = container.add_paragraph()
                _reset_paragraph_format(p)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.left_indent = Pt(14)
                p.add_run("• ")
                __render_markdown_text(p, li.get_text().strip(), style_config)

        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                max_cols = max([len(r.find_all(['td', 'th'])) for r in rows])
                docx_table = container.add_table(rows=0, cols=max_cols)
                try:
                    docx_table.style = 'Table Grid'
                except KeyError:
                    pass
                for i, r in enumerate(rows):
                    cells = r.find_all(['td', 'th'])
                    if cells:
                        row_cells = docx_table.add_row().cells
                        for j, c in enumerate(cells):
                            if j < len(row_cells):
                                target_cell = row_cells[j]
                                p = target_cell.paragraphs[0]
                                _reset_paragraph_format(p)
                                p.text = ""
                                __render_markdown_text(p, c.get_text().strip(), style_config)
                                if c.name == 'th':
                                    for run in p.runs:
                                        run.bold = True

        elif element.name is None:
            text = str(element).strip()
            if text:
                p = container.add_paragraph()
                _reset_paragraph_format(p)
                p.paragraph_format.space_before = Pt(6)
                __render_markdown_text(p, text, style_config)


# ─────────────────────────────────────────────
# HTML → Section Parser
# ─────────────────────────────────────────────

def parse_html_into_sections(payload_html: str) -> dict:
    """
    Parse AI-generated HTML into named sections.
    Returns: { "문제인식": [elements...], "실현가능성": [elements...], ... }
    """
    # Pre-process HTML bold tags into Markdown equivalent
    payload_html = payload_html.replace('<strong>', '**').replace('</strong>', '**')
    payload_html = payload_html.replace('<b>', '**').replace('</b>', '**')

    soup = BeautifulSoup(payload_html, "html.parser")
    all_elements = list(soup.body.children if soup.body else soup.children)

    sections = {}
    current_key = "도입부"
    sections[current_key] = []

    for element in all_elements:
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            text = element.get_text().strip()
            n_key = normalize_section_key(text)
            if n_key:
                current_key = n_key

            if current_key not in sections:
                sections[current_key] = []
            sections[current_key].append(element)
        else:
            if str(element).strip():
                if current_key not in sections:
                    sections[current_key] = []
                sections[current_key].append(element)

    return sections, all_elements


# ─────────────────────────────────────────────
# MODE 1: Coordinate-based Fill (Precision)
# ─────────────────────────────────────────────

def coordinate_based_fill(doc, template_config, payload_html, output_path) -> bool:
    """
    Use hardcoded cell coordinates from the template registry to insert
    AI content into the exact right cells. Maximum accuracy.
    """
    try:
        sections, all_elements = parse_html_into_sections(payload_html)
        style_config = template_config.get("style", {})
        section_map = template_config.get("sections", {})
        mapped_keys = set()

        print(f"📋 Parsed sections: {list(sections.keys())}")
        print(f"📋 Template sections: {list(section_map.keys())}")

        for section_key, coord in section_map.items():
            t_idx = coord["table_index"]
            r_idx = coord["content_row"]
            c_idx = coord["content_col"]

            if t_idx >= len(doc.tables):
                print(f"⚠️ Table index {t_idx} out of range (doc has {len(doc.tables)} tables)")
                continue

            table = doc.tables[t_idx]
            if r_idx >= len(table.rows):
                print(f"⚠️ Row index {r_idx} out of range in table {t_idx}")
                continue

            # Get the unique (deduplicated) cells in the target row
            unique_cells = get_unique_cells(table.rows[r_idx])
            if c_idx >= len(unique_cells):
                print(f"⚠️ Col index {c_idx} out of range in table {t_idx} row {r_idx}")
                continue

            target_cell = unique_cells[c_idx]

            # Find matching section content
            if section_key in sections and len(sections[section_key]) > 0:
                print(f"✅ Inserting '{section_key}' → Table {t_idx}, Row {r_idx}, Col {c_idx}")
                _insert_html_into_cell(target_cell, sections[section_key], style_config)
                mapped_keys.add(section_key)
            else:
                print(f"⚠️ No content found for section '{section_key}'")

        if len(mapped_keys) == 0:
            print("⚠️ Coordinate-based fill failed to map ANY sections. Falling back...")
            return False

        # Handle unmapped sections (append at end)
        unmapped = [k for k in sections.keys() if k not in mapped_keys and k != "도입부"]
        if unmapped:
            print(f"📎 Unmapped sections will be appended: {unmapped}")
            doc.add_page_break()
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("=== 추가 AI 분석 내용 (매핑되지 않음) ===")
            title_run.bold = True
            title_run.font.size = Pt(13)
            for k in unmapped:
                insert_elements_into_container(doc, sections[k], style_config)

        # XML-Level fix: Remove rigid row heights & Floating text overlap
        _fix_table_xml(doc)

        doc.save(output_path)
        print(f"✅ Coordinate-based fill complete → {output_path}")
        return True

    except Exception as e:
        import traceback
        trace_str = traceback.format_exc()
        print(f"❌ Coordinate-based fill error: {e}")
        with open('/tmp/insightflow_error.log', 'w') as f:
            f.write(trace_str)
        return False


# ─────────────────────────────────────────────
# MODE 2: Heuristic Fill (Fallback)
# ─────────────────────────────────────────────

def heuristic_fill(doc, payload_html, output_path) -> bool:
    """
    Fallback for unrecognized templates. Scans table cells for keywords
    and tries to insert content into adjacent cells.
    """
    try:
        sections, all_elements = parse_html_into_sections(payload_html)
        mapped_keys = set()
        
        candidates = {}
        active_n_key = None

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                n_key = normalize_section_key(block.text.strip())
                if n_key and n_key in sections:
                    active_n_key = n_key
                    
            elif isinstance(block, Table):
                table_head_text = ""
                if len(block.rows) > 0:
                    table_head_text = "".join(c.text for c in block.rows[0].cells).replace(" ", "")
                if len(doc.tables) > 1 and ("목차" in table_head_text or "개요" in table_head_text or "요약" in table_head_text):
                    continue

                if active_n_key:
                    largest_cell = None
                    max_area = 0
                    for row in block.rows:
                        unique_cells = []
                        seen_tcs = set()
                        for c in row.cells:
                            if c._tc not in seen_tcs:
                                seen_tcs.add(c._tc)
                                unique_cells.append(c)
                        
                        for c in unique_cells:
                            if len(c.text.strip()) < 300:
                                area = len(row.cells)
                                if area >= max_area:
                                    max_area = area
                                    largest_cell = c
                    
                    if largest_cell:
                        candidates.setdefault(active_n_key, []).append(largest_cell)
                        active_n_key = None

                for r_idx, row in enumerate(block.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        n_key = normalize_section_key(cell_text)

                        if "{{본문}}" in cell_text or "{{내용}}" in cell_text:
                             candidates.setdefault("본문", []).append(cell)
                             continue

                        if n_key and n_key in sections:
                            target_cell = None

                            if r_idx + 1 < len(block.rows):
                                below_row = block.rows[r_idx + 1]
                                if len(set(c._tc for c in below_row.cells)) == 1:
                                    target_cell = below_row.cells[0]

                            if not target_cell and c_idx + 1 < len(row.cells):
                                if row.cells[c_idx + 1]._tc != cell._tc:
                                    target_cell = row.cells[c_idx + 1]

                            if not target_cell and r_idx + 1 < len(block.rows):
                                if block.rows[r_idx + 1].cells[c_idx]._tc != cell._tc:
                                    target_cell = block.rows[r_idx + 1].cells[c_idx]

                            if target_cell:
                                candidates.setdefault(n_key, []).append(target_cell)
                                active_n_key = None
                            else:
                                active_n_key = n_key

        # Now perform the actual insertions into the LAST found cell for each section
        if "본문" in candidates:
            # For exact template tokens, just take the first one
            best_cell = candidates["본문"][0]
            _clear_cell(best_cell)
            insert_elements_into_container(best_cell, all_elements)
            mapped_keys.add("본문")

        for section_key, cell_list in candidates.items():
            if section_key == "본문":
                continue
            # Pick the LAST matched cell - skips the top summary table and targets the detailed body
            best_cell = cell_list[-1]
            _clear_cell(best_cell)
            _insert_html_into_cell(best_cell, sections[section_key])
            mapped_keys.add(section_key)

        # Fallback append
        unmapped = [k for k in sections.keys() if k not in mapped_keys and k != "도입부"]
        if len(mapped_keys) == 0:
            doc.add_page_break()
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("=== AI 초안 본문 (자동 첨부) ===")
            title_run.bold = True
            title_run.font.size = Pt(14)
            insert_elements_into_container(doc, all_elements)
        elif len(unmapped) > 0:
            doc.add_page_break()
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("=== 누락된 AI 구역 내용 (자동 첨부) ===")
            title_run.bold = True
            for k in unmapped:
                insert_elements_into_container(doc, sections[k])

        _fix_table_xml(doc)
        doc.save(output_path)
        print(f"✅ Heuristic fill complete → {output_path}")
        return True

    except Exception as e:
        import traceback
        trace_str = traceback.format_exc()
        print(f"❌ Heuristic fill error: {e}")
        with open('/tmp/insightflow_error.log', 'w') as f:
            f.write(trace_str)
        return False


def _fix_table_xml(doc):
    """Remove rigid trHeight and tblpPr (floating properties) to allow auto-expansion and prevent overlaps."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for table in doc.tables:
        # 1. Un-float table (Prevent text from overlapping behind the table)
        tblPr = table._element.tblPr
        if tblPr is not None:
            for tblpPr in tblPr.findall(f"{ns}tblpPr"):
                tblPr.remove(tblpPr)

        # 2. Prevent cut-off text (Remove fixed row heights)
        for row in table.rows:
            tr = row._tr
            trPr = tr.trPr
            if trPr is not None:
                for trh in trPr.findall(f"{ns}trHeight"):
                    trPr.remove(trh)


# ─────────────────────────────────────────────
# MODE 3: AI Dynamic Fill (처음 보는 양식 전용)
# ─────────────────────────────────────────────

def ai_dynamic_fill(doc, input_path: str, payload_html: str, output_path: str, gemini_api_key: str, announcement_sections: list = None) -> bool:
    """
    AI가 양식 구조를 분석하여 자동으로 섹션 매핑을 생성하고 내용을 삽입.
    처음 보는 모든 양식에 대응 가능.
    """
    try:
        template_config = analyze_template_with_ai(doc, input_path, gemini_api_key, announcement_sections)
        if not template_config:
            print("⚠️ AI 동적 분석 실패 — Heuristic 모드로 전환")
            return False

        print(f"🤖 AI 동적 매핑 사용: {template_config.get('_template_name', 'Unknown')}")

        # AI가 생성한 섹션 매핑의 키 이름을 normalize하여 HTML 파싱 결과와 매칭
        sections, all_elements = parse_html_into_sections(payload_html)
        style_config = template_config.get("style", {})
        section_map = template_config.get("sections", {})
        mapped_keys = set()

        print(f"📋 AI 분석 섹션: {list(section_map.keys())}")
        print(f"📋 HTML 파싱 섹션: {list(sections.keys())}")

        for ai_section_key, coord in section_map.items():
            t_idx = coord.get("table_index", 0)
            r_idx = coord.get("content_row", 0)
            c_idx = coord.get("content_col", 1)

            if t_idx >= len(doc.tables):
                print(f"⚠️ 테이블 {t_idx} 범위 초과")
                continue

            table = doc.tables[t_idx]
            if r_idx >= len(table.rows):
                print(f"⚠️ 행 {r_idx} 범위 초과 (테이블 {t_idx})")
                continue

            unique_cells = get_unique_cells(table.rows[r_idx])
            if c_idx >= len(unique_cells):
                print(f"⚠️ 열 {c_idx} 범위 초과")
                continue

            target_cell = unique_cells[c_idx]

            # AI 섹션 키 → HTML 섹션 키 매핑 시도 (normalize 사용)
            # 1) 직접 매칭
            matched_html_key = None
            if ai_section_key in sections:
                matched_html_key = ai_section_key
            else:
                # 2) normalize로 유사 매칭
                from template_registry import normalize_section_key
                norm_ai = normalize_section_key(ai_section_key)
                for html_key in sections:
                    norm_html = normalize_section_key(html_key)
                    if norm_ai and norm_html and norm_ai == norm_html:
                        matched_html_key = html_key
                        break
                # 3) label로도 시도
                if not matched_html_key:
                    label = coord.get("label", "")
                    norm_label = normalize_section_key(label) if label else None
                    for html_key in sections:
                        norm_html = normalize_section_key(html_key)
                        if norm_label and norm_html and norm_label == norm_html:
                            matched_html_key = html_key
                            break

            if matched_html_key and sections.get(matched_html_key):
                print(f"✅ AI매핑 '{ai_section_key}' → HTML '{matched_html_key}' → Table {t_idx}, Row {r_idx}, Col {c_idx}")
                _insert_html_into_cell(target_cell, sections[matched_html_key], style_config)
                mapped_keys.add(matched_html_key)
            else:
                print(f"⚠️ '{ai_section_key}' 에 매칭되는 HTML 섹션 없음")

        if not mapped_keys:
            print("⚠️ AI 동적 매핑: 매핑된 섹션 없음 → Heuristic 전환")
            return False

        # 미매핑 섹션 추가
        unmapped = [k for k in sections if k not in mapped_keys and k != "도입부"]
        if unmapped:
            print(f"📎 미매핑 섹션 추가: {unmapped}")
            doc.add_page_break()
            p = doc.add_paragraph()
            r = p.add_run("=== 추가 AI 분석 내용 ===")
            r.bold = True
            r.font.size = Pt(13)
            for k in unmapped:
                insert_elements_into_container(doc, sections[k], style_config)

        _fix_table_xml(doc)
        doc.save(output_path)
        print(f"✅ AI 동적 fill 완료 → {output_path}")
        return True

    except Exception as e:
        import traceback
        print(f"❌ AI 동적 fill 오류: {e}")
        traceback.print_exc()
        return False


# ─────────────────────────────────────────────
# Main Entry Point
# ─────────────────────────────────────────────

def fill_docx_template(input_path: str, payload_html: str, output_path: str, gemini_api_key: str = None, announcement_sections: list = None) -> bool:
    """
    마스터 함수: 3단계 전략으로 최적 모드를 자동 선택.

    1단계: 하드코딩 레지스트리 (알려진 양식 → 최고 정밀도)
    2단계: AI 동적 분석 (처음 보는 양식 → Gemini가 구조 파악)
    3단계: 휴리스틱 (AI 불가 시 → 키워드 스캔 폴백)
    """
    try:
        doc = docx.Document(input_path)

        # ── 1단계: 하드코딩 레지스트리 ──
        template_config = detect_template(doc)
        if template_config:
            success = coordinate_based_fill(doc, template_config, payload_html, output_path)
            if success:
                return True
            print("🔄 좌표 모드 실패 → 다음 단계로")
            doc = docx.Document(input_path)

        # ── 2단계: AI 동적 분석 ──
        if gemini_api_key:
            success = ai_dynamic_fill(doc, input_path, payload_html, output_path, gemini_api_key, announcement_sections)
            if success:
                return True
            print("🔄 AI 동적 분석 실패 → Heuristic으로 전환")
            doc = docx.Document(input_path)
        else:
            print("⚠️ Gemini API 키 없음 → AI 동적 분석 스킵")

        # ── 3단계: 휴리스틱 폴백 ──
        return heuristic_fill(doc, payload_html, output_path)

    except Exception as e:
        import traceback
        trace_str = traceback.format_exc()
        print(f"❌ DOCX 처리 오류: {e}")
        with open('/tmp/insightflow_error.log', 'w') as f:
            f.write(trace_str)
        return False
